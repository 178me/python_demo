import shutil
import os
import sys
import json
import logging
import requests
import pypandoc
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from time import sleep
from random import randint, choice
from lxml import etree
from PyQt5.Qt import (QThread)
import re
zhPattern = re.compile(u'[\u4e00-\u9fa5]+')


class MyThread(QThread):
    def __init__(self, func=None, args=None, obj=None, signal=None):
        super(MyThread, self).__init__()
        # 要执行的函数
        self.func = func
        # 函数参数
        self.fun_args = args

    def run(self):
        ''' 执行对应的任务 '''
        result = None
        try:
            if not self.func:
                raise Exception("没有功能被执行")
            if self.fun_args:
                result = self.func(self.fun_args)
            else:
                result = self.func()
        except Exception as e:
            if result:
                logging.info("完美结束啦!")
            else:
                logging.warning("失败了哦")


def to_doc(title, p_list):
    # 新建空白文档
    title = title.replace(".html", "")
    doc1 = Document()
    # 新增文档标题
    doc1.add_heading(title, 0)
    for p in p_list:
        doc1.add_paragraph(p)
    # 保存文件
    doc1.save(f'./无图Doc/{title}.doc')


def to_info():
    error_page = []
    start_page = int(sys.argv[1])
    end_page = int(sys.argv[2])
    for i in range(start_page, end_page):
        try:
            r = requests.get(f"https://www.paidai.com/more.php?page={i}")
            with open(f"./info/info{i}.html", "w", encoding="utf-8") as f:
                f.write(r.text)
            sleep(randint(1, 5))
        except Exception as e:
            logging.exception(e)
            error_page.append(i)
            if len(error_page) > 5:
                break
    print(error_page)
    with open(f"./info/info.txt", "a", encoding="utf-8") as f:
        json.dump(error_page, f)


def to_id_list():
    def sort_by_name(info):
        if info == 'info.txt':
            return -1
        return int(info.replace(".html", "").replace("info", ""))
    info_list = os.listdir("./info")
    info_list.sort(key=sort_by_name, reverse=False)
    id_list = []
    for info in info_list:
        try:
            if info == 'info.txt':
                continue
            html_path = f"./info/{info}"
            print(html_path)
            html = etree.parse(html_path, etree.HTMLParser())
            elems = html.xpath(
                '//*[@class="content-bottom-l-2-m2"]/ul/li/div[2]/span[1]/a')
            for e in elems:
                e_id = f'https:{e.get("href")}'
                print(e_id)
                if "bbs.paidai.com" not in e_id:
                    continue
                #  print(e_id)
                id_list.append({
                    "page": int(info.replace(".html", "").replace("info", "")),
                    "title": e.get("title").replace("?", "？").replace(r"/", "-").replace(":", "：").replace("*", "-").replace("<", "《").replace(">", "》").replace("|", "-").strip(),
                    "url": e_id
                })
            #  break
        except Exception as e:
            logging.exception(e)
            print("出错了")
            break
    #  print(id_list)
    print(len(id_list))
    with open(f"id_list.txt", "w", encoding="utf-8") as f:
        json.dump(id_list, f, indent=4, ensure_ascii=False)


def to_paper_list(page):
    #  print(page)
    #  return
    my_headers = [
        "Mozilla/5.0 (Windows NT 6.3; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.153 Safari/537.36",
        "Mozilla/5.0 (Windows NT 6.1; WOW64; rv:30.0) Gecko/20100101 Firefox/30.0",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_2) AppleWebKit/537.75.14 (KHTML, like Gecko) Version/7.0.3 Safari/537.75.14",
        "Mozilla/5.0 (compatible; MSIE 10.0; Windows NT 6.2; Win64; x64; Trident/6.0)",
        'Mozilla/5.0 (Windows; U; Windows NT 5.1; it; rv:1.8.1.11) Gecko/20071127 Firefox/2.0.0.11',
        'Opera/9.25 (Windows NT 5.1; U; en)',
        'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; .NET CLR 1.1.4322; .NET CLR 2.0.50727)',
        'Mozilla/5.0 (compatible; Konqueror/3.5; Linux) KHTML/3.5.5 (like Gecko) (Kubuntu)',
        'Mozilla/5.0 (X11; U; Linux i686; en-US; rv:1.8.0.12) Gecko/20070731 Ubuntu/dapper-security Firefox/1.5.0.12',
        'Lynx/2.8.5rel.1 libwww-FM/2.14 SSL-MM/1.4.1 GNUTLS/1.2.9',
        "Mozilla/5.0 (X11; Linux i686) AppleWebKit/535.7 (KHTML, like Gecko) Ubuntu/11.04 Chromium/16.0.912.77 Chrome/16.0.912.77 Safari/535.7",
        "Mozilla/5.0 (X11; Ubuntu; Linux i686; rv:10.0) Gecko/20100101 Firefox/10.0 "
    ]
    {
        "User-Agent": "Mozilla/5.0 (X11; Linux i686) AppleWebKit/535.7 (KHTML, like Gecko) Ubuntu/11.04 Chromium/16.0.912.77 Chrome/16.0.912.77 Safari/535.7"
    }
    with open(f"id_list.txt", "r", encoding="utf-8") as f:
        id_list = json.load(f)
    error_page = []
    start_page = page[0]
    end_page = page[1]
    not_found = []
    current_page = start_page
    ip = requests.get("http://localhost:5555/random").text
    proxies = {
        'http': r"http://" + ip,
        'https': r"https://" + ip
    }
    print(proxies)
    i = start_page
    while i < end_page:
        #  for i in range(start_page, end_page):
        info = id_list[i]
        try:
            if os.path.exists(f"./paper/{info['title'].strip()}.html"):
                print("文件已经存在")
                i += 1
                continue
            if id_list[i] == "https:":
                i += 1
                continue
            print(info)
            print(f"开始 {i}:", info["title"].strip())
            headers = {
                "Referer": "https://www.paidai.com/",
                "User-Agent": choice(my_headers)
            }
            print("当前IP", proxies)
            try:
                r = requests.get(info["url"], headers=headers,
                                 proxies=proxies, timeout=(5, 5))
            except:
                ip = requests.get("http://localhost:5555/random").text
                proxies = {
                    'http': r"http://" + ip,
                    'https': r"https://" + ip
                }
                print("IP错误:更换IP地址")
                continue
            if r.status_code != 200:
                ip = requests.get("http://localhost:5555/random").text
                proxies = {
                    'http': r"http://" + ip,
                    'https': r"https://" + ip
                }
                print("请求错误:更换IP地址")
                continue
            with open(f"./paper/{info['title'].strip()}.html", "w", encoding="utf-8") as f:
                f.write(r.text)
            sleep(randint(1, 5))
            i += 1
        except Exception as e:
            logging.exception(e)
            error_page.append(info["url"])
            if len(error_page) > 5:
                break


def to_paper_exits():
    paper_list = os.listdir("./paper")
    #  for paper in paper_list:
    #  print(paper)
    #  name = paper.replace(".html","").strip()
    #  os.rename(f"./paper/{paper}",f"./paper/{name}.html")
    #  return
    with open(f"id_list.txt", "r", encoding="utf-8") as f:
        id_list = json.load(f)
    list1 = []
    for info in id_list:
        title = info["title"].strip() + ".html"
        if title in paper_list:
            continue
        print(title)
        list1.append(info)
    print(len(list1))
    with open(f"./not_exist.txt", "w", encoding="utf-8") as f:
        json.dump(list1, f, indent=4, ensure_ascii=False)


def to_paper_img():
    if os.path.exists("./img_list.txt"):
        with open(f"./img_list.txt", "r", encoding="utf-8") as f:
            list1 = json.load(f)
    else:
        list1 = []
    #  print(len(list1))
    paper_list = os.listdir("./paper")
    img_list = []
    #  print(len(paper_list))
    for img in list1:
        title = img["title"]
        if title in paper_list:
            paper_list.remove(title)
            continue
    #  print(len(paper_list))
    #  return
    for paper in paper_list:
        html_path = f"./paper/{paper}"
        print(html_path)
        html = etree.parse(html_path, etree.HTMLParser())
        elems = html.xpath('//*[@id="topic_content"]/p/img')
        t_l = []
        for e in elems:
            t_l.append(e.get("src"))
        img_list.append({
            "title": paper,
            "img_list": t_l,
        })
        #  break
    list1 += img_list
    list2 = []

    for i, li in enumerate(list1):
        temp = []
        for l in li["img_list"]:
            if not l:
                continue
            if "http:" in l:
                temp.append(l)
        list1[i]["img_list"] = temp
    for li in list1:
        if li["img_list"]:
            list2.append(li)
    #  print(len(list1))
    with open(f"./img_list.txt", "w", encoding="utf-8") as f:
        json.dump(list2, f, indent=4, ensure_ascii=False)


def to_download_img(page):
    #  print(page)
    #  return
    my_headers = [
        "Mozilla/5.0 (Windows NT 6.3; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.153 Safari/537.36",
        "Mozilla/5.0 (Windows NT 6.1; WOW64; rv:30.0) Gecko/20100101 Firefox/30.0",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_2) AppleWebKit/537.75.14 (KHTML, like Gecko) Version/7.0.3 Safari/537.75.14",
        "Mozilla/5.0 (compatible; MSIE 10.0; Windows NT 6.2; Win64; x64; Trident/6.0)",
        'Mozilla/5.0 (Windows; U; Windows NT 5.1; it; rv:1.8.1.11) Gecko/20071127 Firefox/2.0.0.11',
        'Opera/9.25 (Windows NT 5.1; U; en)',
        'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; .NET CLR 1.1.4322; .NET CLR 2.0.50727)',
        'Mozilla/5.0 (compatible; Konqueror/3.5; Linux) KHTML/3.5.5 (like Gecko) (Kubuntu)',
        'Mozilla/5.0 (X11; U; Linux i686; en-US; rv:1.8.0.12) Gecko/20070731 Ubuntu/dapper-security Firefox/1.5.0.12',
        'Lynx/2.8.5rel.1 libwww-FM/2.14 SSL-MM/1.4.1 GNUTLS/1.2.9',
        "Mozilla/5.0 (X11; Linux i686) AppleWebKit/535.7 (KHTML, like Gecko) Ubuntu/11.04 Chromium/16.0.912.77 Chrome/16.0.912.77 Safari/535.7",
        "Mozilla/5.0 (X11; Ubuntu; Linux i686; rv:10.0) Gecko/20100101 Firefox/10.0 "
    ]
    with open(f"img_list.txt", "r", encoding="utf-8") as f:
        img_list = json.load(f)
    error_page = []
    start_page = page[0]
    end_page = page[1]
    ip = requests.get("http://localhost:5555/random").text
    proxies = {
        'http': "http://" + ip,
        'https': "https://" + ip
    }
    print(proxies)
    i = start_page
    while i < end_page:
        #  for i in range(start_page, end_page):
        img = img_list[i]
        try:
            print(img)
            print(f"开始 {i}:", img["title"].strip())
            headers = {
                "User-Agent": choice(my_headers)
            }
            j = 0
            while j < len(img["img_list"]):
                try:
                    filename = img["img_list"][j].replace("?", "？").replace(r"/", "-").replace(":", "：").replace(
                        "*", "-").replace("<", "《").replace(">", "》").replace("|", "-").strip() + ".png"
                    if os.path.exists("./image/" + filename):
                        print("文件已经存在")
                        j += 1
                        continue
                    if "http" not in img["img_list"][j]:
                        print("不是一个url")
                        j += 1
                        continue
                    r = requests.get(
                        img["img_list"][j], headers=headers, proxies=proxies, timeout=(5, 5))
                    print("当前IP", proxies["https"])
                    if r.status_code != 200:
                        ip = requests.get("http://localhost:5555/random").text
                        proxies = {
                            'http': "http://" + ip,
                            'https': "https://" + ip
                        }
                        print("请求错误:更换IP地址")
                        continue
                    with open(f"./image/{filename}", "wb") as f:
                        f.write(r.content)
                    j += 1
                except Exception as e:
                    ip = requests.get("http://localhost:5555/random").text
                    proxies = {
                        'http': "http://" + ip,
                        'https': "https://" + ip
                    }
                    print(e)
                    print("IP错误:更换IP地址")
                    continue
            i += 1
        except Exception as e:
            logging.exception(e)
            error_page.append(img["url"])
            if len(error_page) > 5:
                break


def html_to_docx(title_text, html_path, docx_path):
    html = etree.parse(html_path, etree.HTMLParser())
    # html
    for e in html.xpath("/*/*/*"):
        if e.tag == "div" and e.get("id") == "container":
            continue
        e.getparent().remove(e)
    for e in html.xpath('//*[@id="container"]/*'):
        if e.tag == "div" and e.get("id") == "container_body":
            continue
        e.getparent().remove(e)
    for e in html.xpath('//*[@id="container_body"]/*'):
        if e.tag == "div" and e.get("id") == "container_l":
            continue
        e.getparent().remove(e)
    for e in html.xpath('//*[@id="container_l"]/*'):
        if e.tag == "div" and e.get("class") == "content_box":
            continue
        e.getparent().remove(e)
    for e in html.xpath('//*[@class="content_box"]/*'):
        if e.tag == "div" and e.get("id") == "topic_content":
            continue
        e.getparent().remove(e)
    head = html.xpath("//head")[0]
    title = etree.Element("title")
    title.text = title_text
    head.append(title)
    save_html_path = f"./cache/{title_text}.html"
    html.write(save_html_path)
    pypandoc.convert_file(save_html_path, 'docx', outputfile=docx_path)


def to_run(_range=[0, 100]):
    print(f"开始转换：{_range[0]} - {_range[1]}")
    root = "./网页文章"
    root2 = "./Doc文章"
    list1 = os.listdir(root)
    i = _range[0]
    end = _range[1]
    while i < end:
        print(f"开始第{i}个:{list1[i]}")
        title_text = list1[i].replace("%", "百分比")
        if title_text[0] == ".":
            title_text = title_text[1:]
        title_text = title_text.replace(".html", "")
        if os.path.exists(f"{root2}/{title_text}.docx"):
            print(f"文件已存在:{root2}/{title_text}.docx")
            i += 1
            continue
        try:
            html_to_docx(
                title_text, f"{root}/{list1[i]}", f"{root2}/{title_text}.docx")
        except Exception as e:
            print(e)
            continue
        i += 1


def main():
    with open(f"id_list.txt", "r", encoding="utf-8") as f:
        id_list = json.load(f)
    print(len(id_list))
    off = 500
    start = 0
    end = off
    list1 = []
    for i in range(64):
        #  print(start, end)
        list1.append([start, end])
        start += off
        end += off
        if end > len(id_list):
            list1.append([start, len(id_list)])
            break
    print(len(list1))
    #  return
    thread_list = []
    for li in list1:
        thread_list.append(MyThread(func=to_paper_list, args=li))
    for thread in thread_list:
        thread.start()
    while True:
        print("主线程等待中")
        sleep(60)


def main2():
    with open(f"img_list.txt", "r", encoding="utf-8") as f:
        id_list = json.load(f)
    print(len(id_list))
    off = 50
    start = 0
    end = off
    list1 = []
    for i in range(32):
        #  print(start, end)
        list1.append([start, end])
        start += off
        end += off
        if end > len(id_list):
            list1.append([start, len(id_list)])
            break
    print(len(list1))
    #  return
    thread_list = []
    for li in list1:
        thread_list.append(MyThread(func=to_download_img, args=li))
    for thread in thread_list:
        thread.start()
    while True:
        print("主线程等待中")
        sleep(60)


def main3():
    root = "./网页文章"
    list1 = os.listdir(root)
    off = 100
    start = 0
    end = start + off
    list2 = []
    for i in range(15):
        #  print(start, end)
        list2.append([start, end])
        start += off
        end = start + off
        if end > len(list1):
            list2.append([start, len(list1)])
            break
    print(list2)
    print(f"线程数量:{len(list2)}")
    sleep(3)
    thread_list = []
    for li in list2:
        thread_list.append(MyThread(func=to_run, args=li))
    for thread in thread_list:
        thread.start()
        sleep(0.1)
    while True:
        for i, thread in enumerate(thread_list):
            if thread.isRunning():
                print(f"转换中：{thread.fun_args[0]} - {thread.fun_args[1]}")
            else:
                start += off
                end = start + off
                if end > len(list1):
                    end = len(list1)
                    thread.fun_args = [start, end]
                    thread.start()
                    break
                thread.fun_args = [start, end]
                thread.start()
        if end >= len(list1):
            break
        sleep(1)
    while True:
        for i, thread in enumerate(thread_list):
            if thread.isRunning():
                print(f"转换中：{thread.fun_args[0]} - {thread.fun_args[1]}")
            else:
                print(f"完成转换：{thread.fun_args[0]} - {thread.fun_args[1]}")
        sleep(60)


def move_file():
    list1 = os.listdir("./paper")
    for paper in list1:
        html_path = f"./paper/{paper}"
        print(html_path)
        html = etree.parse(html_path, etree.HTMLParser())
        elems = html.xpath('//*[@id="topic_content"]/p//img')
        if not elems:
            shutil.move(html_path, f"./无图文章/{paper}")
            continue
        shutil.move(html_path, f"./有图文章/{paper}")


def test():
    to_run()


#  main()
#  to_id_list()
#  to_paper_exits()
#  to_paper_name()
#  to_paper_img()
#  to_paper_list([1, 5000])
#  to_download_img([0, 1000])
#  main2()
#  test()
main3()
